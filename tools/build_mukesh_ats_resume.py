from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


OUT_DIR = Path("/Users/mac/Desktop/university_pwc")
DOCX_PATH = OUT_DIR / "Mukesh_Thokala_Flutter_Developer_ATS_Resume.docx"
PDF_PATH = OUT_DIR / "Mukesh_Thokala_Flutter_Developer_ATS_Resume.pdf"


resume = {
    "name": "Thokala Mukesh",
    "title": "Flutter Developer | Mobile App Stability | API Integration | Firebase",
    "contact": [
        "+91-8341257923",
        "mukeshthokala4@gmail.com",
        "linkedin.com/in/mukesh-t-00a012285",
        "github.com/Thokalamukesh",
    ],
    "summary": (
        "Flutter Developer with 1+ year of experience building production-grade cross-platform "
        "mobile applications using Flutter, Dart, Firebase, Firestore, and REST APIs. Strong in "
        "API integration, mobile app lifecycle, state management, debugging, performance optimization, "
        "responsive UI/UX flows, and Android deployment readiness. Comfortable owning production issues, "
        "collaborating with backend teams, and improving app stability, responsiveness, and reliability "
        "in fast-paced product environments."
    ),
    "skills": [
        "Mobile: Flutter, Dart, Android Studio, cross-platform mobile development, responsive UI implementation",
        "App Stability & Performance: crash analysis fundamentals, debugging, ANR investigation, memory leak awareness, UI rendering optimization, app lifecycle, load-time improvement",
        "APIs & Data: REST API integration, JSON handling, response validation, error handling, data caching, backend connectivity, GraphQL fundamentals",
        "Firebase: Firebase Authentication, Cloud Firestore, real-time data synchronization, Firebase services",
        "State & Architecture: state management, reusable UI components, clean UI architecture, cart/order workflows",
        "Release & Collaboration: Git, GitHub, version control discipline, Android build testing, release readiness, production issue resolution",
        "Languages & Tools: Java, Python, SQL, HTML, CSS, WordPress, Visual Studio Code",
    ],
    "experience": [
        {
            "company": "Sirixo",
            "role": "Flutter Developer",
            "date": "1 Year",
            "bullets": [
                "Developed and maintained production-grade Flutter applications using reusable UI components and scalable feature structure.",
                "Integrated REST APIs with response validation, error handling, and data caching to improve reliability across mobile workflows.",
                "Improved app stability and responsiveness by debugging UI issues, optimizing navigation flows, and reducing avoidable runtime failures.",
                "Implemented USB thermal printer integrations with status callback handling for real-time transaction confirmation.",
                "Engineered kiosk auto-navigation and self-ordering workflows optimized for touch-first tablet and kiosk deployments.",
                "Supported release readiness through Android build testing, Git-based version control, issue triage, and production bug fixes.",
            ],
        },
        {
            "company": "KranviK - Technology and Corporate",
            "role": "Flutter Developer Intern",
            "date": "March 2024 - October 2024",
            "bullets": [
                "Built responsive cross-platform mobile screens in Flutter and Dart with optimized UI rendering and clean component structure.",
                "Integrated REST APIs and Firebase services for authentication, real-time data updates, and backend-driven app flows.",
                "Collaborated with senior developers and backend teams to refine UI/UX flows, resolve API issues, and improve data handling.",
                "Performed debugging, issue resolution, and stability testing across Android builds to support production-ready deployment.",
                "Improved app performance by identifying rendering issues, validating edge cases, and tightening mobile user flows.",
            ],
        },
    ],
    "projects": [
        {
            "name": "Kiosk Self-Order System",
            "tech": "Flutter, Dart, REST APIs, State Management, USB Thermal Printer",
            "bullets": [
                "Developed a cross-platform self-ordering POS application for kiosk and tablet devices, allowing customers to browse products, customize orders, and complete checkout flows independently.",
                "Designed a responsive, touch-friendly UI with smooth navigation, cart management, real-time price calculation, and structured order confirmation.",
                "Integrated REST APIs for dynamic product listing and order submission, with optimized state handling for large product catalogs.",
                "Implemented transaction confirmation and printer status handling to improve reliability in production-style kiosk environments.",
            ],
        },
        {
            "name": "E-Commerce Mobile App",
            "tech": "Flutter, Firebase Authentication, Firestore",
            "bullets": [
                "Built a mobile e-commerce app with product listings, cart management, wishlist, checkout flow, and order tracking.",
                "Integrated Firebase Authentication and Firestore for secure user management and real-time product data synchronization.",
                "Executed Android build testing, debugging, and performance checks to improve app responsiveness and reliability.",
            ],
        },
        {
            "name": "Blood Group Detection Using Fingerprint",
            "tech": "Machine Learning, Image Processing, Python",
            "bullets": [
                "Designed a prototype machine learning workflow for fingerprint feature extraction and pattern recognition.",
                "Applied image processing techniques to support classification workflows for healthcare diagnostic research.",
            ],
        },
    ],
    "education": [
        "B.Tech in Computer Science and Information Technology, Siddharth Institute of Engineering and Technology, Puttur | GPA: 7.8 | Expected 2025"
    ],
    "certifications": [
        "Full Stack Web Development - TechCiti Software Consulting Pvt Ltd",
        "Flutter Assessment - LearnTube by CarrierNinja",
        "Industrial and Internet of Things 4.0 - NPTEL",
    ],
}


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.03

    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(2)

    styles["Heading 1"].font.size = Pt(10.6)
    styles["Heading 2"].font.size = Pt(9.6)


def add_bottom_border(paragraph) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run(title.upper())
    run.bold = True
    add_bottom_border(p)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(1.6)
    p.paragraph_format.line_spacing = 1.02
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.0)


def build_docx() -> None:
    doc = Document()
    set_document_styles(doc)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(0)
    r = name.add_run(resume["name"])
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(17)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1.5)
    r = title.add_run(resume["title"])
    r.font.name = "Arial"
    r.font.size = Pt(9.5)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    r = contact.add_run(" | ".join(resume["contact"]))
    r.font.name = "Arial"
    r.font.size = Pt(8.8)

    add_section(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.04
    r = p.add_run(resume["summary"])
    r.font.name = "Arial"
    r.font.size = Pt(9.1)

    add_section(doc, "Technical Skills")
    for item in resume["skills"]:
        add_bullet(doc, item)

    add_section(doc, "Professional Experience")
    for job in resume["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0.8)
        r = p.add_run(f'{job["role"]} | {job["company"]}')
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9.4)
        r2 = p.add_run(f' | {job["date"]}')
        r2.font.name = "Arial"
        r2.font.size = Pt(9.1)
        for bullet in job["bullets"]:
            add_bullet(doc, bullet)

    add_section(doc, "Projects")
    for project in resume["projects"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1.5)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(project["name"])
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9.4)
        r2 = p.add_run(f' | {project["tech"]}')
        r2.italic = True
        r2.font.name = "Arial"
        r2.font.size = Pt(8.9)
        for bullet in project["bullets"]:
            add_bullet(doc, bullet)

    add_section(doc, "Education")
    for item in resume["education"]:
        add_bullet(doc, item)

    add_section(doc, "Certifications")
    for item in resume["certifications"]:
        add_bullet(doc, item)

    doc.save(DOCX_PATH)


def build_pdf() -> None:
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.9,
        leading=10.1,
        spaceAfter=2.2,
        alignment=TA_LEFT,
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=12,
        firstLineIndent=-7,
        bulletIndent=0,
        spaceAfter=1.9,
    )
    section = ParagraphStyle(
        "Section",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=10.2,
        leading=11,
        spaceBefore=5,
        spaceAfter=2,
        borderWidth=0.55,
        borderColor=colors.black,
        borderPadding=1,
    )
    job_style = ParagraphStyle(
        "Job",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=9.2,
        leading=10.3,
        spaceBefore=2,
        spaceAfter=0.6,
    )
    name_style = ParagraphStyle(
        "Name",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=18,
        alignment=TA_CENTER,
        spaceAfter=0,
    )
    center = ParagraphStyle(
        "Center",
        parent=body,
        alignment=TA_CENTER,
        fontSize=8.7,
        leading=9.5,
        spaceAfter=2,
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
    )
    story = [
        Paragraph(resume["name"], name_style),
        Paragraph(resume["title"], center),
        Paragraph(" | ".join(resume["contact"]), center),
        Spacer(1, 1),
        Paragraph("PROFESSIONAL SUMMARY", section),
        Paragraph(resume["summary"], body),
        Paragraph("TECHNICAL SKILLS", section),
    ]
    for item in resume["skills"]:
        story.append(Paragraph(item, bullet, bulletText="-"))

    story.append(Paragraph("PROFESSIONAL EXPERIENCE", section))
    for job in resume["experience"]:
        story.append(Paragraph(f'<b>{job["role"]} | {job["company"]}</b> | {job["date"]}', job_style))
        for item in job["bullets"]:
            story.append(Paragraph(item, bullet, bulletText="-"))

    story.append(Paragraph("PROJECTS", section))
    for project in resume["projects"]:
        story.append(Paragraph(f'<b>{project["name"]}</b> | <i>{project["tech"]}</i>', job_style))
        for item in project["bullets"]:
            story.append(Paragraph(item, bullet, bulletText="-"))

    story.append(Paragraph("EDUCATION", section))
    for item in resume["education"]:
        story.append(Paragraph(item, bullet, bulletText="-"))

    story.append(Paragraph("CERTIFICATIONS", section))
    for item in resume["certifications"]:
        story.append(Paragraph(item, bullet, bulletText="-"))

    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
